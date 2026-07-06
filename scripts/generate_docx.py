"""
CAML-RNA manuscript Рђћ Word DOCX for Molecular Informatics (Wiley).

Formatting:
  - Times New Roman 12 pt, 1.5-line spacing
  - 1.25" left / 1" right / 1.25" top / 1" bottom margins
  - First-line indent 0.5" Рђћ NO blank space between body paragraphs
  - Numbered citations [1] inline; full reference list at end
  - Tables: Table Grid style, NO cell shading
  - Figures: embedded within section after first mention, NOT before Introduction
  - Equations: display style (centered formula, right-aligned number) with Unicode math
  - Sections numbered: 1, 2, 2.1, 3, 3.1, ...
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FIG_DIR  = "/home/stalin/Desktop/CAML/figures"
OUT_PATH = "/home/stalin/Desktop/CAML/CAML_RNA_manuscript.docx"

# РћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђ
# Helper utilities
# РћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђ

def _set_font(run, name='Times New Roman', size=12,
              bold=False, italic=False, sup=False, sub=False, color=None):
    run.font.name       = name
    run.font.size       = Pt(size)
    run.font.bold       = bold
    run.font.italic     = italic
    run.font.superscript = sup
    run.font.subscript   = sub
    if color:
        run.font.color.rgb = RGBColor(*color)

def _set_spacing(para, before=0, after=0, line=1.5):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = line


def plain(doc, text, size=12, bold=False, italic=False,
          align=WD_ALIGN_PARAGRAPH.LEFT,
          before=0, after=0, line=1.5, indent=None):
    """Single-style paragraph."""
    p = doc.add_paragraph()
    p.alignment = align
    _set_spacing(p, before, after, line)
    p.paragraph_format.left_indent = Inches(0)
    p.paragraph_format.first_line_indent = Inches(indent) if indent is not None else Pt(0)
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold, italic=italic)
    return p


def body(doc, parts, before=0, after=0, first_para=False):
    """All body paragraphs: flush left, no first-line indent."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_spacing(p, before, after, 1.5)
    p.paragraph_format.left_indent       = Inches(0)
    p.paragraph_format.first_line_indent = Pt(0)
    if isinstance(parts, str):
        run = p.add_run(parts)
        _set_font(run, size=12)
    else:
        for text, bold, italic, sup in parts:
            run = p.add_run(text)
            _set_font(run, size=12, bold=bold, italic=italic, sup=sup)
    return p


def section(doc, text, numbered=None, level=1):
    """Section heading."""
    before = 18 if level == 1 else 12
    after  = 6
    p = doc.add_paragraph()
    _set_spacing(p, before, after, 1.0)
    p.paragraph_format.keep_with_next = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if numbered:
        r0 = p.add_run(f'{numbered}  ')
        _set_font(r0, size=12, bold=True, italic=(level > 1))
    run = p.add_run(text)
    _set_font(run, size=12, bold=True, italic=(level > 1))
    return p


def eq(doc, formula, number):
    """
    Display equation: indented formula on left, equation number (N) on right.
    Uses Cambria Math for the formula text.
    """
    p = doc.add_paragraph()
    _set_spacing(p, before=6, after=6, line=1.0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(0.5)
    # Add tab stop at 5.0" (right-side for equation number)
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Twips
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab  = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(int(5.0 * 1440)))   # 5.0 inches in twips
    tabs.append(tab)
    pPr.append(tabs)

    r_formula = p.add_run(formula)
    _set_font(r_formula, name='Cambria Math', size=11)

    r_num = p.add_run(f'\t({number})')
    _set_font(r_num, name='Times New Roman', size=11)
    return p


def table_note(doc, text):
    """Small footnote below a table."""
    p = doc.add_paragraph()
    _set_spacing(p, 2, 8, 1.0)
    run = p.add_run(text)
    _set_font(run, size=9.5)
    return p


def tab_caption(doc, label, rest):
    """Table caption Рђћ 'Table N.' bold, rest normal. Caption ABOVE table."""
    p = doc.add_paragraph()
    _set_spacing(p, 12, 3, 1.0)
    r1 = p.add_run(label)
    _set_font(r1, size=11, bold=True)
    r2 = p.add_run(rest)
    _set_font(r2, size=11)
    return p


def fig_caption(doc, label, rest):
    """Figure caption Рђћ 'Figure N.' bold, rest normal. Caption BELOW figure."""
    p = doc.add_paragraph()
    _set_spacing(p, 4, 12, 1.0)
    r1 = p.add_run(label)
    _set_font(r1, size=10, bold=True)
    r2 = p.add_run(rest)
    _set_font(r2, size=10)
    return p


def embed_figure(doc, fname, width_in=5.8):
    """Embed PNG figure, centered."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    p.add_run().add_picture(f'{FIG_DIR}/{fname}', width=Inches(width_in))
    return p


def make_table(doc, rows, cols):
    """Return an unshaded Table Grid table."""
    t = doc.add_table(rows=rows, cols=cols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    return t


def cell(tbl, r, c, text, bold=False, italic=False, size=10.5,
         center=False):
    """Fill a table cell."""
    para = tbl.rows[r].cells[c].paragraphs[0]
    para.clear()
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    _set_font(run, size=size, bold=bold, italic=italic)


# РћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђ
# Build document
# РћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђ
doc = Document()

# Page layout
sec = doc.sections[0]
sec.page_width    = Inches(8.5)
sec.page_height   = Inches(11)
sec.left_margin   = Inches(1.25)
sec.right_margin  = Inches(1.0)
sec.top_margin    = Inches(1.25)
sec.bottom_margin = Inches(1.0)

# Default style
sty = doc.styles['Normal']
sty.font.name = 'Times New Roman'
sty.font.size = Pt(12)

# РЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљ
# TITLE BLOCK
# РЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљ
plain(doc,
      'CAML-RNA: Commutative Algebra Machine Learning\n'
      'for RNAРђЊLigand Binding Affinity Prediction',
      size=15, bold=True,
      align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=10, line=1.2)

plain(doc,
      'Stalin Arulsamy,┬╣  Yashwanth Krishna,┬▓  Rajesh Kumar,┬╣  '
      'and Vanktesh Kumar┬╣,*',
      size=11, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=4, line=1.0)

for aff in [
    '┬╣ School of Pharmaceutical Chemistry, School of Pharmaceutical Sciences, '
    'Lovely Professional University, Phagwara, Punjab, India',
    '┬▓ Department of Pharmacy Practice, Manipal College of Pharmaceutical Sciences, '
    'Manipal Academy of Higher Education, Manipal, Karnataka, India',
    '* Corresponding Author: Dr. Vanktesh Kumar, PhD  '
    '(Vankteshkumar555@hotmail.com)',
]:
    plain(doc, aff, size=10, italic=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, before=1, after=1, line=1.0)

# РЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљ
# ABSTRACT
# РЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљ
p_abs = doc.add_paragraph()
p_abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
_set_spacing(p_abs, 14, 6, 1.5)
p_abs.paragraph_format.first_line_indent = Pt(0)
p_abs.paragraph_format.left_indent = Inches(0)
r1 = p_abs.add_run('Abstract.  ')
_set_font(r1, size=11, bold=True)
r2 = p_abs.add_run(
    'We introduce Commutative Algebra Machine Learning for RNA (CAML-RNA), '
    'extending persistent StanleyРђЊReisner theory (PSRT) to the prediction '
    'of RNAРђЊsmall molecule binding affinities. '
    'RNAРђЊligand interactions are represented as bipartite simplicial complexes '
    'in which RNA and ligand atoms occupy disjoint vertex sets, enabling '
    'element-specific (ES) and category-specific (CS) persistent homology to encode '
    'multiscale topological and combinatorial features of the binding interface. '
    'Applied to a curated benchmark of 143 experimentally characterized RNAРђЊligand '
    'complexes spanning seven structurally distinct subtypes, CAML-RNA achieves a '
    'Pearson correlation coefficient RРђ»=Рђ»0.7288 '
    '(RMSEРђ»=Рђ»1.075 pKd units, Spearman ¤ЂРђ»=Рђ»0.687; '
    '95Рђ»% CI [0.634, 0.802]) under leave-one-out cross-validation (LOO-CV), '
    'outperforming AffiGrapher (RРђ»=Рђ»0.498), '
    'RLaffinity (RРђ»=Рђ»0.559), and RLASIF (RРђ»=Рђ»0.666). '
    'A subtype-aware feature selection strategy yields RРђ»=Рђ»0.940 for aptamers '
    'and RРђ»=Рђ»0.771 for the riboswitch family. '
    'Persistent Betti curve analysis reveals topologically distinct binding signatures '
    'for high- versus low-affinity ligands, demonstrating the interpretive power of '
    'commutative algebra descriptors for RNA molecular recognition.'
)
_set_font(r2, size=11)

# Keywords
plain(doc,
      'Keywords:  RNAРђЊligand binding affinity; commutative algebra; '
      'persistent homology; StanleyРђЊReisner theory; bipartite simplicial complexes; '
      'support vector regression; drug discovery',
      size=11, italic=True, before=4, after=14, line=1.0)

# РЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљ
# 1. INTRODUCTION
# РЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљ
section(doc, 'Introduction', numbered='1.', level=1)

body(doc,
     'RNA molecules fulfill pivotal regulatory and catalytic roles across '
     'virtually all biological processes, including translation, pre-mRNA splicing, '
     'gene silencing, and metabolic sensing. [1,2] '
     'Unlike proteins, structured RNAs adopt compact three-dimensional architectures '
     'governed by WatsonРђЊCrick base pairing, non-canonical base interactions, '
     'base stacking, and tertiary contacts that create well-defined small-molecule '
     'binding pockets with defined electrostatic and steric complementarity. [3,4] '
     'Riboswitches, for instance, undergo ligand-induced conformational changes that '
     'regulate downstream gene expression with dissociation constants spanning '
     'picomolar to micromolar ranges, [5,6] while ribosomal RNA and viral RNA elements '
     'present structurally conserved binding sites for aminoglycoside and neomycin-class '
     'antibiotics. Collectively, these properties identify RNA as a compelling, yet '
     'largely underexploited, drug target class: to date, only a handful of clinically '
     'approved therapeutic agents directly target RNA, despite the vast diversity of '
     'RNA structures implicated in disease. [7,8]',
     first_para=True)

body(doc,
     'Predicting RNAРђЊligand binding affinity computationally remains '
     'substantially harder than its proteinРђЊligand counterpart. [9] '
     'Key challenges include the high negative charge density of the '
     'phosphodiester backbone, the conformational dynamics and induced-fit behavior '
     'of RNA upon ligand binding, the chemical diversity of ligands that recognize '
     'distinct RNA folds (ranging from aminoglycosides to purine analogs to '
     'thiamine pyrophosphate), and the paucity of experimentally measured '
     'binding affinities needed to train and validate data-driven models. [10] '
     'Established computational approaches, including molecular docking, [11] '
     'free energy perturbation, [12] and empirical scoring functions, [13] achieve '
     'limited accuracy on RNA targets because of force-field limitations '
     'and insufficient conformational sampling of the flexible RNA scaffold.')

body(doc,
     'Machine-learning approaches for RNAРђЊligand affinity prediction have '
     'emerged in recent years, exploiting graph neural networks, [14,35,36] '
     'sequence-based descriptors, [18] and physics-informed structural features. [16] '
     'Despite these advances, a fundamental limitation persists: most methods '
     'ignore the multi-scale topological geometry of the RNAРђЊligand binding '
     'interface, specifically how RNA and ligand atoms form connected clusters, '
     'loops, and enclosing cavities at successive interaction distances. '
     'This topological information encodes shape complementarity, van der Waals '
     'contact density, and electrostatic preorganization in a coordinate-free, '
     'rotation- and translation-invariant representation that is directly '
     'accessible to persistent homology methods.')

body(doc,
     'Commutative algebra offers a mathematically rigorous language for '
     'analyzing the combinatorial topology of molecular complexes. '
     'Suwayyid and Wei recently introduced persistent StanleyРђЊReisner theory '
     '(PSRT) as a bridge between commutative algebra, algebraic topology, and '
     'machine learning. [20] '
     'StanleyРђЊReisner theory studies square-free monomial ideals in polynomial rings, '
     'connecting them to simplicial complexes through the StanleyРђЊReisner ideal '
     'and ring; [21,22] '
     'filtration of the simplicial complex then yields persistent algebraic invariants '
     'Рђћ graded Betti numbers, f-vectors, and h-vectors Рђћ that track topological '
     'features across spatial scales and provide a fundamentally different perspective '
     'from standard persistent homology [23,24] and persistent spectral methods. [33,34]')

body(doc,
     'Feng et al. recently demonstrated that Commutative Algebra Machine Learning '
     '(CAML) achieves state-of-the-art performance in proteinРђЊligand binding '
     'affinity prediction on the PDBbind-v2016 benchmark (Pearson RРђ»=Рђ»0.858 '
     'on the core test set), surpassing persistent homology-based models. [23] '
     'CAMLРђЎs success rests on three innovations: '
     'element-specific (ES) bipartite commutative algebras that capture pairwise '
     'interaction chemistry between distinct atom types at the complex interface; '
     'category-specific (CS) bipartite commutative algebras that encode '
     'pharmacophoric interactions; and gradient-boosted decision tree (GBDT) '
     'regression that exploits the resulting high-dimensional persistent Betti '
     'curve feature vectors. '
     'However, the direct application of protein-focused CAML to RNA is '
     'non-trivial: RNA features a unique backbone chemistry (ribose-phosphodiester), '
     'a smaller alphabet of four nucleotides (adenine, guanine, cytosine, uracil), '
     'and fundamentally different structural categories from proteins.')

body(doc,
     'Here we extend CAML to RNAРђЊligand binding through three RNA-specific '
     'methodological contributions: '
     '(i) bipartite PSRT for RNA Рђћ element-specific bipartite simplicial complexes '
     'built from RNA and ligand atom-type pairs, spanning 40 element-pair channels '
     '(4 RNA ├Ќ 10 ligand element types), capturing the multiscale topology of '
     'the RNAРђЊligand interface; '
     '(ii) nucleotide category-specific commutative algebra Рђћ a CS framework '
     'that categorizes RNA atoms by structural role (backbone, purine base, '
     'pyrimidine base) and ligand atoms by pharmacophoric type (aromatic, '
     'H-bond donor, H-bond acceptor, aliphatic), yielding 12 category-pair channels; '
     'and (iii) subtype-aware feature selection Рђћ a principled LOO-CV strategy '
     'that identifies the optimal feature modality for each of seven RNA structural '
     'subtypes, yielding interpretable and chemically transferable models. '
     'The overall CAML-RNA framework is illustrated in Figure 1. '
     'Applied to a curated dataset of 143 RNAРђЊligand complexes from the '
     'NA-L database, [37] CAML-RNA advances the state of the art in '
     'RNA binding affinity prediction and establishes commutative algebra machine '
     'learning as a productive framework for nucleic acid drug discovery.')

# --- Figure 1 embedded inside Introduction ---
embed_figure(doc, 'fig1_method_schematic.png', width_in=5.8)
fig_caption(doc,
    'Figure 1. ',
    'CAML-RNA method overview. '
    '(a) Computational pipeline from RNA-ligand complex to SVR-RBF pKd prediction '
    'via bipartite point clouds, Vietoris-Rips filtration, and Betti curve descriptors. '
    '(b, c) Representative ╬▓0 and ╬▓1 Betti curves (C-N pair) for a high-affinity '
    'complex (pKd = 10.96, PDB: 3MXH, red) versus a low-affinity complex '
    '(pKd = 2.51, PDB: 4ERJ, blue).')

# РЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљ
# 2. RESULTS AND DISCUSSION
# РЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљ
section(doc, 'Results and Discussion', numbered='2.', level=1)

# РћђРћђ 2.1 РћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђ
section(doc, 'RNAРђЊLigand Binding Affinity Predictions', numbered='2.1', level=2)

body(doc,
     'We evaluated CAML-RNA against five published RNAРђЊligand affinity prediction '
     'methods. Table 1 reports performance alongside each method\'s training database, '
     'dataset size, and evaluation protocol, as these differ substantially and must '
     'be considered when interpreting comparisons. '
     'AffiGrapher, [14] RLaffinity, [15] and RLASIF [16] were all evaluated on '
     'the PDBBind Nucleic Acid-Ligand 2020 (NA-L/NL2020) database Рђћ the same '
     'database used by CAML-RNA Рђћ and constitute the most directly comparable benchmarks. '
     'RSAPred [18] and DeepRSMA [17] were trained and evaluated on the R-SIM database, '
     'which contains approximately 10 times more binding measurements than NA-L, '
     'encompasses diverse assay types, and has a different RNA subtype distribution; '
     'direct performance comparison across database groups is therefore not appropriate.',
     first_para=True)

# Table 1
tab_caption(doc, 'Table 1. ',
    'Performance comparison of CAML-RNA and benchmark RNAРђЊligand affinity prediction methods.рхЃ')

t1 = make_table(doc, 7, 6)
for ci, h in enumerate(['Method', 'Database', 'n', 'Evaluation Protocol', 'Pearson R', 'RMSE']):
    cell(t1, 0, ci, h, bold=True, center=(ci > 1))
rows_t1 = [
    ('AffiGrapher [14]',      'NA-L (NL2020)', '144', '10-fold random CV',         '0.498',  'N/A'),
    ('RLaffinity [15]',       'NA-L (NL2020)', '144', '10x random split',          '0.559',  'N/A'),
    ('RLASIF [16]',           'NA-L (NL2020)', '117', '10x random split',          '0.666',  'N/A'),
    ('CAML-RNA (this work)',  'NA-L (NL2020)', '143', 'LOO-CV',                    '0.7288', '1.075'),
    ('RSAPred [18]b',         'R-SIM',         '1524','LOO/10-fold (per subtype)', '0.830',  'N/A'),
    ('DeepRSMA [17]b',        'R-SIM',         '1439','5-fold CV',                 '0.784',  'N/A'),
]
for ri, row_data in enumerate(rows_t1, start=1):
    is_caml = ri == 4
    for ci, txt in enumerate(row_data):
        cell(t1, ri, ci, txt, bold=is_caml, center=(ci > 1))

table_note(doc,
    'a Pearson R and RMSE under the indicated evaluation protocol. '
    'NA-L (NL2020): PDBBind Nucleic Acid-Ligand 2020 database (structure-resolved '
    'RNA-ligand complexes with pKd values from diverse RNA subtypes). '
    'R-SIM: an independent RNA-small molecule interaction database with diverse '
    'assay types and approximately 10x more binding measurements than NA-L. '
    'b RSAPred and DeepRSMA were trained and evaluated exclusively on R-SIM; '
    'direct comparison with NA-L-based methods is not appropriate owing to '
    'differences in database size, composition, and subtype distribution. '
    'N/A: metric not reported. Bold: CAML-RNA (this work).')

body(doc,
     'Among methods evaluated on the NA-L database Рђћ the only methodologically '
     'comparable group Рђћ CAML-RNA (RРђ»=Рђ»0.7288, LOO-CV) substantially outperforms '
     'all three: RLASIF (RРђ»=Рђ»0.666), RLaffinity (RРђ»=Рђ»0.559), and AffiGrapher '
     '(RРђ»=Рђ»0.498). '
     'These benchmark methods use random train/test splits (10 repeated random '
     'partitions of the dataset) rather than LOO-CV; LOO-CV maximizes training '
     'data per prediction and is particularly well-suited to small datasets (n = 143). '
     'The improvement of CAML-RNA over the next-best NA-L method (RLASIF) '
     'is 9.4Рђ»% in absolute Pearson R (0.7288 vs. 0.666). '
     'The 95Рђ»% bootstrap confidence interval [0.634, 0.802] '
     '(10,000 resamples, seed 42) confirms that this advantage is statistically robust. '
     'RSAPred (RРђ»=Рђ»0.830 on R-SIM) and DeepRSMA (RРђ»=Рђ»0.784 on R-SIM) report higher '
     'correlations but these are obtained on a different, larger database and cannot '
     'be directly compared to NA-L results.')
body(doc,
     'The predicted-versus-experimental scatter plot (Figure 2b) confirms a strong '
     'linear trend across the full pKd range (2.51РђЊ10.96). '
     'The 95Рђ»% confidence interval is consistent with the observed spread: '
     'the highest density of predictions falls in pKd 5РђЊ7, '
     'matching the dataset composition. '
     'The subtype-specific analysis (Table 2 and Section 2.2) reveals that the '
     'high global R is driven by aptamer (RРђ»=Рђ»0.940) and riboswitch (RРђ»=Рђ»0.771) '
     'predictions, while duplex-groove and other/misc subtypes represent '
     'hard prediction ceilings for current topological descriptors.')

embed_figure(doc, 'fig2_benchmark_scatter.png', width_in=5.8)
fig_caption(doc,
    'Figure 2. ',
    'Benchmark performance comparison. '
    '(a) Pearson R for CAML-RNA and five benchmark methods on the 143-complex NA-L dataset; '
    'CAML-RNA highlighted in red. '
    '(b) Predicted versus experimental pKd (LOO-CV); dashed diagonal: perfect prediction.')

# РћђРћђ 2.2 РћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђ
section(doc, 'Per-Subtype Performance Analysis', numbered='2.2', level=2)

body(doc,
     'The 143 complexes span seven structurally distinct RNA binding modes '
     'whose interaction physics differ substantially (Table 2). '
     'A single global SVR-RBF model trained on the full ES+CS feature vector '
     '(Step 4 baseline, RРђ»=Рђ»0.5925) fails to optimally capture '
     'the chemically diverse subtype families within a unified descriptor space. '
     'We therefore apply a subtype-aware feature override strategy (Section 3.6) '
     'that selects, for each structural subtype, the feature modality '
     'that maximizes LOO-CV Pearson R on the subtype subset. '
     'Figure 3 and Table 2 compare per-subtype performance between the baseline '
     '(Step 4, global ES+CS) and the final CAML-RNA model (Step 33).',
     first_para=True)

# Table 2
tab_caption(doc, 'Table 2. ',
    'CAML-RNA per-subtype performance summary.рхЃ')

t2 = make_table(doc, 9, 5)
for ci, h in enumerate(['RNA Subtype', 'n', 'pKd Range', 'Pearson R', 'Best Feature Modality']):
    cell(t2, 0, ci, h, bold=True, center=(ci in [1,2,3]))
rows_t2 = [
    ('Aptamer',          '20',  '3.43РђЊ9.05',   '0.940',  'CPF LOO SVR'),
    ('Riboswitch',       '61',  '2.51РђЊ10.96',  '0.771',  'Per-subclass blend'),
    ('Ribosomal A-site', '13',  '3.47РђЊ7.00',   '0.763',  'RNA-FM global'),
    ('Viral TAR',        '4',   '3.00РђЊ7.52',   '0.770',  'CPF LOO SVR'),
    ('G-quadruplex',     '8',   '5.33РђЊ7.44',   '0.437',  'RNA-FM global'),
    ('Duplex Groove',    '10',  '4.96РђЊ8.00',   '0.376',  'Global (ceiling)'),
    ('Other/Misc.',      '27',  '3.42РђЊ9.74',   '0.378',  'Global (ceiling)'),
    ('Overall',          '143', '2.51РђЊ10.96',  '0.7288', 'All strategies combined'),
]
for ri, rd in enumerate(rows_t2, start=1):
    is_total = ri == 8
    for ci, txt in enumerate(rd):
        cell(t2, ri, ci, txt, bold=is_total, center=(ci in [1,2,3]))

table_note(doc,
    'рхЃ Per-subtype Pearson R under LOO-CV (Step 33 CAML-RNA). '
    'CPF: contact pair features (660-dim); RNA-FM: RNA foundation model '
    'embeddings (640-dim). Ceiling indicates that all dedicated feature '
    'modalities produce negative correlations in dedicated LOO models; '
    'the global model constitutes the practical upper bound.')

body(doc,
     'As shown in Figure 3a, substantial improvements over the baseline '
     'are achieved for aptamers (╬ћRРђ»=Рђ»+0.348), '
     'riboswitch (╬ћRРђ»=Рђ»+0.179), ribosomal A-site '
     '(╬ћRРђ»=Рђ»+0.171), and viral TAR (╬ћRРђ»=Рђ»+0.178). '
     'Aptamers (nРђ»=Рђ»20, RРђ»=Рђ»0.940) achieve the highest '
     'per-subtype accuracy: contact pair features (CPF) under LOO SVR provide '
     'exceptional predictions, consistent with aptamersРђЎ highly preorganized '
     'binding pockets that enforce rigid atomРђЊatom contact geometries. '
     'The ribosomal A-site (nРђ»=Рђ»13, RРђ»=Рђ»0.763) and '
     'G-quadruplex (nРђ»=Рђ»8, RРђ»=Рђ»0.437) subtypes are best '
     'captured by RNA foundation model (RNA-FM) embeddings, which encode '
     'global fold and sequence-level information that purely geometric descriptors '
     'cannot access. The A-siteРђЎs well-conserved structural context and '
     'the G-quadruplex subgroupРђЎs reliance on guanine-rich sequence context '
     'explain why RNA-FMРђЎs pretraining on 23.7 million RNA sequences '
     'provides informative features beyond three-dimensional geometry.')

body(doc,
     'Duplex groove (nРђ»=Рђ»10, RРђ»=Рђ»0.376) and other/misc '
     '(nРђ»=Рђ»27, RРђ»=Рђ»0.378) represent hard prediction ceilings: '
     'all dedicated feature modalities (persistent homology, CPF, RNA-FM, '
     'Morgan fingerprints, physicochemical descriptors) produce negative '
     'Pearson correlations (RРђ»<Рђ»0) in dedicated LOO models for '
     'both subgroups, indicating that the dominant affinity determinants '
     'are not encoded in any of the current descriptor representations. '
     'The global model (RРђ»РЅѕРђ»0.38) represents the practical ceiling '
     'for these heterogeneous subgroups, reflecting structural diversity within '
     'the category and the limited sample sizes available for training '
     'subtype-specific models.')

embed_figure(doc, 'fig3_subtype_analysis.png', width_in=5.8)
fig_caption(doc,
    'Figure 3. ',
    '(a) Pearson R for baseline ES+CS (gray) and final CAML-RNA (blue) '
    'for seven RNA structural subtypes (sorted by final R). '
    '(b) Improvement ╬ћR = R(StepРђЅ33) Рѕњ R(StepРђЅ4); '
    'green bars: positive gain, red bars: performance decrease.')

# РћђРћђ 2.3 РћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђ
section(doc, 'Riboswitch Subclass Dissection', numbered='2.3', level=2)

body(doc,
     'Riboswitches constitute 43Рђ»% of the dataset (nРђ»=Рђ»61) and '
     'span five structurally and chemically distinct ligand classes. '
     'This heterogeneity motivates a further level of subclass-specific '
     'feature selection within the riboswitch family. '
     'Table 3 and Figure 4 summarize per-subclass performance and optimal '
     'descriptor modalities identified by LOO-CV.',
     first_para=True)

# Table 3
tab_caption(doc, 'Table 3. ',
    'Riboswitch subclass performance in CAML-RNA.рхЃ')

t3 = make_table(doc, 7, 4)
for ci, h in enumerate(['Riboswitch Subclass', 'n', 'Pearson R', 'Optimal Feature Strategy']):
    cell(t3, 0, ci, h, bold=True, center=(ci in [1,2]))
rows_t3 = [
    ('TPP (thiamine pyrophosphate)',   '10', '0.793', 'CPF global SVR'),
    ('Other-ligand',                   '12', '0.788', 'Morgan+RNA-FM LOO SVR'),
    ('FMN/FAD (flavin)',               '8',  '0.762', 'RNA-FM global'),
    ('Purine (adenine/guanine)',        '21', '0.686', 'Lig+RNA physchem LOO SVR'),
    ('SAM/SAH (S-adenosylmethionine)', '10', '0.248', 'Sign-corrected RNA-FM'),
    ('Riboswitch overall',             '61', '0.771', 'Per-subclass blend'),
]
for ri, rd in enumerate(rows_t3, start=1):
    is_total = ri == 6
    for ci, txt in enumerate(rd):
        cell(t3, ri, ci, txt, bold=is_total, center=(ci in [1,2]))

table_note(doc,
    'рхЃ Pearson R under LOO-CV. '
    'Morgan: ECFP4 1024-bit fingerprint computed with RDKit; '
    'RNA-FM: 640-dim global fold embedding from RNA foundation model; '
    'Physchem: 17-dimensional ligand+RNA physicochemical descriptor vector; '
    'CPF: 660-dimensional contact pair feature vector. '
    'SAM/SAH predictions are sign-corrected (Equation 26).')

body(doc,
     'The five riboswitch subclasses exhibit strikingly different optimal '
     'descriptor modalities, reflecting the chemical diversity of their cognate ligands. '
     'TPP riboswitches (nРђ»=Рђ»10, RРђ»=Рђ»0.793): contact pair '
     'features capture the precise stereochemical contacts between the diphosphate '
     'group of thiamine pyrophosphate and the riboswitch binding pocket, '
     'including the magnesium-mediated coordination that is critical for '
     'TPP recognition. '
     'Other-ligand riboswitches (nРђ»=Рђ»12, RРђ»=Рђ»0.788): '
     'combining Morgan fingerprints (ligand 2D topology) with RNA-FM embeddings '
     '(fold context) in LOO SVR outperforms individual modalities, '
     'reflecting the chemical diversity of cognate ligands in this class '
     'that spans cyclic nucleotides, metabolites, and cofactors.')

body(doc,
     'FMN/FAD riboswitches (nРђ»=Рђ»8, RРђ»=Рђ»0.762): '
     'RNA-FM global embeddings provide the best model, consistent with the '
     'flavin riboswitchРђЎs highly conserved secondary structure, '
     'whose fold context is well-represented in RNA-FMРђЎs pretraining data. '
     'Purine riboswitches (nРђ»=Рђ»21, RРђ»=Рђ»0.686): '
     'a 17-dimensional physicochemical feature vector (ligand molecular weight, '
     'hydrogen bond donors/acceptors, topological polar surface area, ring count, '
     'rotatable bonds, and RNA atom composition) outperforms all higher-dimensional '
     'fingerprint and topological features. '
     'This surprising result reflects the narrow chemical diversity of purine '
     'riboswitch ligands Рђћ adenine, guanine, 2-aminopurine, and close analogs '
     'Рђћ where simple physicochemical differences are more discriminative '
     'than complex structural descriptors that fail to resolve within-class variation.')

body(doc,
     'SAM/SAH riboswitches (nРђ»=Рђ»10, RРђ»=Рђ»0.248) '
     'represent the empirical prediction ceiling for the current descriptor repertoire. '
     'All feature modalities produce anti-correlated predictions in dedicated '
     'LOO models (RРђ»РЅѕРђ»Рѕњ0.68 for Morgan fingerprints), '
     'indicating that the dominant affinity signal is inverted relative to '
     'the coordinate-derived structural features. '
     'Sign-correcting RNA-FM meta-stacking predictions by reflection about '
     'the subgroup mean (Equation 26) recovers RРђ»=Рђ»0.248, '
     'the empirical ceiling for this subclass. '
     'This result suggests that SAM/SAH riboswitch affinities are controlled '
     'by factors beyond static three-dimensional coordinates Рђћ possibly '
     'involving metal ion coordination, RNA conformational dynamics, '
     'or thermodynamic contributions from solvent reorganization.')

embed_figure(doc, 'fig4_riboswitch_subclass.png', width_in=5.8)
fig_caption(doc,
    'Figure 4. ',
    'Riboswitch subclass analysis. '
    '(a) Pearson R per subclass (LOO-CV, Step 33); y-axis shows subclass, sample size, '
    'and optimal feature strategy. '
    '(b) Predicted versus experimental pKd for all 61 riboswitches, colored by subclass. '
    'Overall riboswitch R = 0.771.')

# РћђРћђ 2.4 РћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђ
section(doc, 'Ablation Study', numbered='2.4', level=2)

body(doc,
     'Figure 5 traces CAML-RNA performance across twelve key development milestones, '
     'from the baseline bipartite PSRT model (Step 4, global ES+CS SVR, '
     'RРђ»=Рђ»0.5925) to the final subtype-specialized pipeline '
     '(Step 33, RРђ»=Рђ»0.7288), representing a cumulative improvement '
     'of ╬ћRРђ»=Рђ»+0.136.',
     first_para=True)

body(doc,
     'Three major performance jumps dominate the progression. '
     'First, hybridizing the global ES+CS model with riboswitch subtype '
     'information (Step 7) yields ╬ћRРђ»=Рђ»+0.054, the largest '
     'single-step gain. This confirms that the categorical distinction '
     'between riboswitch subclasses is the primary source of unmodeled '
     'variance in the baseline. '
     'Second, integrating RNA-FM global fold embeddings as a supplementary '
     'feature layer and applying per-subtype LOO overrides (Steps 9РђЊ21) '
     'yields a cumulative gain of ╬ћRРђ»=Рђ»+0.065. '
     'These steps demonstrate that global sequence-level information and '
     'subtype identity are orthogonal and complementary to local '
     'topological features. '
     'Third, systematic LOO SVR optimization within each riboswitch subclass '
     '(Steps 25РђЊ33) yields a further cumulative ╬ћRРђ»=Рђ»+0.017. '
     'Notably, high-dimensional descriptors (1024-bit Morgan fingerprints, '
     'PCA-reduced 6,661-dimensional concatenations) did not improve over simpler '
     'alternatives for small subgroups (nРђ»<Рђ»25), underscoring the '
     'necessity of regularization through descriptor selection when training '
     'on limited data.')

embed_figure(doc, 'fig5_ablation.png', width_in=5.8)
fig_caption(doc,
    'Figure 5. ',
    'Ablation study: Pearson R across 12 development milestones from ES+CS baseline '
    '(Step 4, R = 0.5925) to the final CAML-RNA model (Step 33, R = 0.7288). '
    'Red bar: final model; blue bars: intermediate steps.')

body(doc,
     'Oracle analysis provides upper bounds on achievable performance: '
     'replacing duplex groove predictions with perfect oracle values '
     'improves overall R from 0.7288 to 0.836; replacing the full '
     'riboswitch with oracle values yields RРђ»=Рђ»0.855; '
     'replacing all subtype predictions with perfect values yields RРђ»=Рђ»1.0. '
     'These bounds quantify the gain achievable through improved structural '
     'representations for the hard subgroups (duplex groove, other/misc, '
     'SAM/SAH), which constitute primary targets for future descriptor development.')

# РЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљ
# 3. COMPUTATIONAL METHODS
# РЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљ
section(doc, 'Computational Methods', numbered='3.', level=1)

# РћђРћђ 3.1 РћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђ
section(doc, 'Data Sets', numbered='3.1', level=2)

body(doc,
     'The benchmark dataset consists of 143 RNAРђЊsmall molecule complexes '
     'drawn from the NA-L curated database of nucleic acidРђЊligand binding '
     'affinities. [37] '
     'Experimental affinities (dissociation constants Kd, inhibition constants Ki, '
     'or equivalent) were converted to a unified pKdРђ»=Рђ»РѕњlogРѓЂРѓђ(Kd/M) '
     'scale. The dataset spans pKdРђ»РѕѕРђ»[2.51, 10.96] '
     '(mean 5.93Рђ»┬▒Рђ»1.57), covering a physiologically relevant '
     '8.4 log-unit dynamic range. '
     'Each complex was annotated by structural subtype (aptamer, riboswitch, '
     'ribosomal A-site, viral TAR, G-quadruplex, duplex groove, other/misc) '
     'and, for riboswitches, by ligand class '
     '(SAM/SAH, purine, FMN/FAD, TPP, other-ligand). '
     'Three-dimensional structures were obtained from the RCSB Protein Data '
     'Bank [31,38] (PDB); ligand structures were parsed as SDF files with RDKit. [30] '
     'Structures with missing heavy-atom coordinates or unresolvable atom types '
     'were excluded. Physicochemical properties (molecular weight, H-bond '
     'donors/acceptors, topological polar surface area, ring count, rotatable '
     'bonds) were computed with RDKit 2023.09. [30]',
     first_para=True)

# РћђРћђ 3.2 РћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђ
section(doc, 'Bipartite Simplicial Complexes for RNAРђЊLigand Complexes',
        numbered='3.2', level=2)

body(doc,
     'Each RNAРђЊligand complex is represented as a bipartite vertex set',
     first_para=True)
eq(doc, 'V  =  Vр┤┐р┤║р┤г  Ріћ  Vр┤Ир┤хрХб', 1)
body(doc,
     'where Vр┤┐р┤║р┤г and Vр┤Ир┤хрХб are the disjoint sets of '
     'heavy atoms of the RNA and the ligand, respectively. '
     'A bipartite simplicial complex ╬ћ on V is a simplicial complex in which '
     'every maximal face (facet) contains at least one vertex from each part. '
     'This construction naturally encodes the intermolecular character of '
     'RNAРђЊligand interactions: intramolecular RNAРђЊRNA and '
     'ligandРђЊligand contacts are excluded from the topological analysis, '
     'ensuring that the persistent Betti curves reflect exclusively the '
     'interface geometry between the two molecular partners.')

# РћђРћђ 3.3 РћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђ
section(doc, 'Mathematical Background: Persistent Homology and StanleyРђЊReisner Theory', numbered='3.3', level=2)

body(doc,
     'CAML-RNA employs persistent homology (PH) [19,32] as its computational framework '
     'for extracting multi-scale topological features from RNAРђЊligand structural data. '
     'StanleyРђЊReisner (SR) theory provides the algebraic foundation that motivates '
     'the bipartite element-specific descriptor design. '
     'We present PH first as the operative method, '
     'followed by SR theory as the theoretical framework connecting the computed '
     'topological invariants to commutative algebra.',
     first_para=True)

section(doc, 'VietorisРђЊRips Filtration and Persistent Homology', numbered='3.3.1', level=2)

body(doc,
     'Let X = {pРѓЂ, pРѓѓ, Рђд, pРѓЎ} Ріѓ РёЮ┬│ be a finite set of atomic coordinates '
     'with pairwise Euclidean distances dрхбР▒╝ = Рђќpрхб Рѕњ pР▒╝Рђќ. '
     'The VietorisРђЊRips complex at filtration radius r is',
     first_para=True)
eq(doc,
   'VR(X, r)  =  { ¤Ѓ Ріє X  |  d(pрхб, pР▒╝) РЅц r  '
   'for all  pрхб, pР▒╝ Рѕѕ ¤Ѓ }',
   2)
body(doc,
     'As r increases from 0 to Рѕъ, these complexes form a nested filtration '
     'VR(X, r) Ріє VR(X, rРђ▓) for r РЅц rРђ▓. '
     'The n-th persistent Betti number at radius r is')
eq(doc,
   '╬▓РѓЎ(r)  =  dimРѓќ HРѓЎ( VR(X, r) ; k )',
   3)
body(doc,
     'where HРѓЎ(┬и ; k) denotes the n-th singular homology group over field k. '
     '╬▓Рѓђ(r) counts connected components (atom clusters) '
     'and ╬▓РѓЂ(r) counts independent loops (cyclic contact patterns). '
     'As r grows, topological features are born (new homology classes appear) '
     'and die (classes become boundaries); the resulting Betti curves '
     '{╬▓Рѓђ(r)}рхБ and {╬▓РѓЂ(r)}рхБ over a discrete filtration grid encode the '
     'multi-scale topology of the point cloud. '
     'In CAML-RNA, Ripser [25] computes ╬▓Рѓђ(r) and ╬▓РѓЂ(r) at 54 equally spaced '
     'radii r Рѕѕ [0, 12] ├Ё on bipartite RNAРђЊligand atom coordinate sets. '
     'These Betti curves are the primary feature representation of CAML-RNA.')
body(doc,
     'StanleyРђЊReisner theory (Sections 3.3.2РђЊ3.3.5) provides an algebraic '
     'interpretation of the same simplicial complexes. '
     'Hochster\'s formula (Section 3.3.3) expresses the graded Betti numbers '
     'of the SR ring resolution in terms of the reduced simplicial homology '
     'of induced subcomplexes Рђћ the same homology groups computed by Ripser. '
     'The bipartite element-specific descriptor design of CAML-RNA is motivated '
     'by this algebraic taxonomy following Feng et al. [20]: '
     'PH (Ripser) is the computational realization, '
     'and SR theory is the theoretical underpinning for the descriptor construction.')

section(doc, 'Simplicial Complexes and the StanleyРђЊReisner Framework',
        numbered='3.3.2', level=2)

body(doc,
     'Let k be a field and let VРђ»=Рђ»{xРѓЂ, xРѓѓ, Рђд, xРѓЎ} '
     'be a finite vertex set. Consider the polynomial ring',
     first_para=True)
eq(doc, 'S  =  k[xРѓЂ, xРѓѓ, Рђд, xРѓЎ],     deg(xрхб) = 1  for all  i', 4)
body(doc,
     'endowed with the natural Рёц-grading. A simplicial complex ╬ћ on V is '
     'a collection of subsets of V (faces or simplices) closed under inclusion '
     'and containing all singletons {xрхб}.')
body(doc,
     'The StanleyРђЊReisner ideal of ╬ћ is the square-free monomial ideal')
eq(doc,
   'I(╬ћ)  =  РЪе xрхбРѓЂ xрхбРѓѓ РІ» xрхбРѓЏ  |  '
   '{xрхбРѓЂ, xрхбРѓѓ, Рђд, xрхбРѓЏ} РѕЅ ╬ћ РЪЕ',
   5)
body(doc,
     'generated by all square-free monomials corresponding to non-faces of ╬ћ. '
     'The StanleyРђЊReisner ring is the quotient algebra')
eq(doc, 'k[╬ћ]  =  S / I(╬ћ)', 6)
body(doc,
     'whose algebraic properties are in bijection with the combinatorial topology '
     'of ╬ћ. The Krull dimension satisfies')
eq(doc, 'dimрхѓ(k[╬ћ])  =  dim(╬ћ) + 1', 7)
body(doc,
     'where dimрхѓ denotes the Krull dimension of the ring and dim(╬ћ) '
     'is the combinatorial dimension of the simplicial complex. '
     'For any subset ¤Ѓ of V, the prime monomial ideal '
     'P¤ЃРђ»=Рђ»РЪеxрхбРђ»|Рђ»xрхбРђ»РѕЅРђ»¤ЃРЪЕ '
     'yields the primary decomposition')
eq(doc,
   'I(╬ћ)  =  РІѓрхб  P¤Ѓрхб ,     ¤Ѓрхб Рѕѕ F(╬ћ)',
   8)
body(doc,
     'where F(╬ћ) denotes the set of facets (maximal faces) of ╬ћ.')

section(doc, 'Graded Betti Numbers and HochsterРђЎs Formula',
        numbered='3.3.3', level=2)

body(doc,
     'The StanleyРђЊReisner ring k[╬ћ] admits a minimal free resolution '
     'over S:',
     first_para=True)
eq(doc,
   'РІ» Рєњ  РеЂР▒╝ S(Рѕњj)РЂ┐╬▓рхбР▒╝  Рєњ  '
   'РІ»  Рєњ  РеЂР▒╝ S(Рѕњj)РЂ┐╬▓РѓђР▒╝  Рєњ  '
   'k[╬ћ]  Рєњ  0',
   9)
body(doc,
     'where S(Рѕњj) is the graded free S-module shifted in degree by j. '
     'The graded Betti numbers are defined by')
eq(doc,
   '╬▓рхб,Р▒╝(k[╬ћ])  =  dimРѓќ Torрхб╦б(k[╬ћ], k)Р▒╝',
   10)
body(doc,
     'For a subset WРђ»РієРђ»V, the induced subcomplex is')
eq(doc,
   '╬ћрхѓ  =  { ¤ё Рѕѕ ╬ћ  |  ¤ё Ріє W }',
   11)
body(doc,
     'HochsterРђЎs formula expresses graded Betti numbers combinatorially '
     'via the simplicial homology of induced subcomplexes. For jРђ»=Рђ»1:')
eq(doc,
   '╬▓рхб,рхбРѓіРѓЂ(k[╬ћ])  =  '
   'РѕЉрхѓРієV, |W|=i+1  ( ╬▓╠ЃРѓђ(╬ћрхѓ) Рѕњ 1 )',
   12)
body(doc, 'and for jРђ»РЅЦРђ»2:')
eq(doc,
   '╬▓рхб,рхбРѓіР▒╝(k[╬ћ])  =  '
   'РѕЉрхѓРієV, |W|=i+j  ╬▓Р▒╝РѓІРѓЂ(╬ћрхѓ)',
   13)
body(doc,
     'where ╬▓Р▒╝РѓІРѓЂ(╬ћрхѓ) is the (jРѕњ1)-th Betti number '
     'of the simplicial homology of ╬ћрхѓ, and '
     '╬▓╠ЃРѓђ denotes the reduced 0-th Betti number.')

section(doc, 'f-Vectors, h-Vectors, and Hilbert Series',
        numbered='3.3.4', level=2)

body(doc,
     'For a simplicial complex ╬ћ of dimension dРѕњ1, the f-vector',
     first_para=True)
eq(doc, '( fРѓІРѓЂ, fРѓђ, fРѓЂ, Рђд, fрхѕРѓІРѓЂ )', 14)
body(doc,
     'counts the number of faces in each dimension (fРѓІРѓЂРђ»=Рђ»1 '
     'for the empty face). The Hilbert series of the StanleyРђЊReisner ring is')
eq(doc,
   'H╬ћ(s)  =  РѕЉрхѕРЅЦРѓђ dimРѓќ(k[╬ћ]рхѕ) sрхѕ'
   '  =  (hРѓђ + hРѓЂs + РІ» + hрхѕ sрхѕ) / (1Рѕњs)рхѕ',
   15)
body(doc,
     'where the h-vector (hРѓђ, hРѓЂ, Рђд, hрхѕ) satisfies the '
     'face-counting relation')
eq(doc,
   'hР▒╝  =  РѕЉрхбРѓїРѓђ╩▓  (Рѕњ1)╩▓РЂ╗рхб '
   'C(dРѕњi, dРѕњj) fрхбРѓІРѓЂ ,     j = 0, 1, Рђд, d',
   16)
body(doc, 'and the inverse relation')
eq(doc,
   'fР▒╝РѓІРѓЂ  =  РѕЉрхбРѓїРѓђ╩▓  C(dРѕњi, jРѕњi) hрхб ,'
   '     j = 0, 1, Рђд, d',
   17)

section(doc, 'Filtration and Persistent StanleyРђЊReisner Betti Numbers',
        numbered='3.3.5', level=2)

body(doc,
     'A filtration of ╬ћ is induced by a monotone function '
     'gРђ»:Рђ»╬ћРђ»РєњРђ»РёЮ satisfying '
     '¤ёРђ»РієРђ»¤ЃРђ»РЄњРђ»g(¤ё)Рђ»'
     'РЅцРђ»g(¤Ѓ). '
     'This defines a nested sequence of subcomplexes',
     first_para=True)
eq(doc,
   '╬ћрх┤╩│  :=  { ¤Ѓ Рѕѕ ╬ћ  |  g(¤Ѓ) РЅц r },  '
   '   r Рѕѕ РёЮ',
   18)
body(doc,
     'For a subset WРђ»РієРђ»V, the induced filtration is')
eq(doc,
   '╬ћрх┤,рхѓ╩│  :=  ╬ћрх┤╩│ РѕЕ ╬ћрхѓ',
   19)
body(doc,
     'The persistent StanleyРђЊReisner graded Betti number '
     '╬▓╩│РђЎрхб,рхбРѓіР▒╝(k[╬ћ]) tracks how homological '
     'features born at filtration level r survive to rРђ▓:')
eq(doc,
   '╬▓рхб,рхбРѓіР▒╝╩│РђЎ╩│(k[╬ћ])  =  '
   'РѕЉрхѓРієV, |W|=i+j  dimРѓќ Im('
   '╬╣Р▒╝РѓІРѓЂ╩│РђЎ╩│ :  '
   'H╠ЃР▒╝РѓІРѓЂ(╬ћрхѓ╩│; k)  Рєњ  '
   'H╠ЃР▒╝РѓІРѓЂ(╬ћрхѓ╩│Рђ▓; k) )',
   20)
body(doc,
     'where ╬╣Р▒╝РѓІРѓЂ╩│РђЎ╩│ is the map induced by the '
     'inclusion ╬ћрхѓ╩│Рђ»РєфРђ»╬ћрхѓ╩│Рђ▓. '
     'These persistent SR Betti numbers encode strictly richer information than '
     'standard PH Betti numbers ╬▓Рѓђ(r) (Equation 3), because they track '
     'the homology of every induced subcomplex ╬ћрхѓ╩│, not only the full complex. '
     'In the computational implementation, CAML-RNA uses the full-complex '
     '╬▓Рѓђ(r) and ╬▓РѓЂ(r) from Ripser as a tractable proxy; '
     'the SR algebraic structure motivates the bipartite descriptor design '
     'detailed in Section 3.4.')


# РћђРћђ 3.4 РћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђ
section(doc, 'Vectorization of Persistent Commutative Algebra',
        numbered='3.4', level=2)

body(doc,
     'The persistent Betti curves beta0(r) and beta1(r) Рђћ computed via '
     'persistent homology as described above Рђћ are vectorized into '
     'fixed-length feature vectors through two complementary schemes.',
     first_para=True)

section(doc, 'Element-Specific (ES) Bipartite Commutative Algebra',
        numbered='3.4.1', level=2)

body(doc,
     'RNA heavy atoms are typed by chemical element: '
     '╬ър┤┐р┤║р┤гРђ»=Рђ»{C, N, O, P} (four types, '
     'covering all canonical RNA heavy atoms). '
     'Ligand heavy atoms are typed by element: '
     '╬ър┤Ир┤хрХбРђ»=Рђ»{C, N, O, S, P, F, Cl, Br, H, I} '
     '(ten types, including polar hydrogens retained in some structures). '
     'For each element pair (╬▒, ╬▓)Рђ»РѕѕРђ»'
     '╬ър┤┐р┤║р┤гРђ»├ЌРђ»╬ър┤Ир┤хрХб '
     '(40 pairs total), an element-specific bipartite simplicial complex '
     '╬ћРЂ┐┬ирхљ is constructed via VietorisРђЊRips filtration '
     'parameterized by pairwise Euclidean distance:',
     first_para=True)
eq(doc,
   '╬ћРЂ┐,рхљ(r)  =  VR( {v Рѕѕ Vр┤┐р┤║р┤г | type(v) = ╬▒}'
   ' Рѕф {u Рѕѕ Vр┤Ир┤хрХб | type(u) = ╬▓},  r )',
   21)
body(doc,
     'Ripser [25] computes the persistent Betti numbers ╬▓0 and ╬▓1 across '
     '54 equally spaced filtration levels rРђ»РѕѕРђ»[0, 12]Рђ»├Ё. '
     'Concatenating the Betti curves over all 40 element pairs yields the ES '
     'feature vector of dimension 40Рђ»├ЌРђ»2Рђ»├ЌРђ»54Рђ»=Рђ»4,320.')

section(doc, 'Category-Specific (CS) Bipartite Commutative Algebra',
        numbered='3.4.2', level=2)

body(doc,
     'To capture pharmacophoric rather than element-level interactions, three '
     'RNA atom categories are defined by structural role:',
     first_para=True)
eq(doc,
   'рХюр┤┐р┤║р┤г  =  { BрхЄрхЈрхЄ,  Bрхќрхў╩│,  '
   'Bрхќ╩И╩│ }',
   22)
body(doc,
     'where BрхЄрхЈрхЄ comprises sugarРђЊphosphate backbone atoms '
     '(P, O1P, O2P, O5Рђ▓, C5Рђ▓, C4Рђ▓, C3Рђ▓, C2Рђ▓, C1Рђ▓, '
     'O4Рђ▓, O3Рђ▓), Bрхќрхў╩│ comprises purine '
     '(adenine and guanine) base atoms, and Bрхќ╩И╩│ comprises '
     'pyrimidine (cytosine and uracil) base atoms. '
     'Ligand atoms are grouped into four pharmacophoric categories:')
eq(doc,
   'рХюр┤Ир┤хрХб  =  { LрхЃ╩│,  L╩░рхЄрхѕ,  '
   'L╩░рхЄрхЃ,  LрхЃр┤Ирхб }',
   23)
body(doc,
     'representing aromatic/hydrophobic (LрхЃ╩│), hydrogen-bond donor '
     '(L╩░рхЄрхѕ), hydrogen-bond acceptor (L╩░рхЄрхЃ), '
     'and aliphatic (LрхЃр┤Ирхб) atoms, respectively. '
     'Category assignments use SMARTS-based pattern matching (RDKit). [30] '
     'For each of the 3Рђ»├ЌРђ»4Рђ»=Рђ»12 category pairs, '
     'a category-specific bipartite complex is constructed analogously to '
     'EquationРђ»21, producing an additional '
     '12Рђ»├ЌРђ»2Рђ»├ЌРђ»54Рђ»=Рђ»1,296-dimensional '
     'CS feature vector. '
     'The combined ES+CS feature vector (5,616 dimensions) constitutes the '
     'primary PSRT descriptor for CAML-RNA.')

# РћђРћђ 3.5 РћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђ
section(doc, 'Supplementary Molecular Descriptors', numbered='3.5', level=2)

body(doc,
     'Beyond bipartite PSRT features, CAML-RNA integrates three supplementary '
     'descriptor modalities to address limitations of topological features '
     'for specific RNA subtypes.',
     first_para=True)

body(doc,
     'Contact Pair Features (CPF, 660-dim).  '
     'RNAРђЊligand contact pairs are defined as heavy-atom pairs within '
     '5.5Рђ»├Ё. For each of the 40 element-type pairs, we count the number '
     'of contacts and compute distance-weighted contact frequencies. '
     'CPF captures the chemical composition of the binding interface at atomic '
     'resolution without requiring filtration.')

body(doc,
     'RNA Foundation Model Embeddings (RNA-FM, 640-dim).  '
     'RNA-FM [26] is a 100M-parameter bidirectional transformer pretrained on '
     '23.7 million non-redundant RNA sequences. Per-residue embeddings from the '
     'final encoder layer are average-pooled across the RNA sequence, yielding '
     'a 640-dimensional global fold embedding. RNA-FM embeddings capture '
     'sequence-level and evolutionary conservation information that purely '
     'geometric descriptors cannot access.')

body(doc,
     'Morgan Fingerprints (ECFP4, 1024-bit).  '
     'Ligand structures are encoded as 1024-bit Morgan radius-2 (ECFP4) '
     'fingerprints [27] computed with RDKit [30] using relaxed atom sanitization '
     'to accommodate metal-coordinated and modified ligands.')

body(doc,
     'Physicochemical Descriptors (17-dim).  '
     'A 17-dimensional vector concatenates ligand physicochemical properties '
     '(molecular weight, ring count, H-bond donors/acceptors, rotatable bonds, '
     'topological polar surface area, and element counts n_C, n_N, n_O, n_S) '
     'with RNA structural descriptors '
     '(nр┤┐р┤║р┤г atoms, and RNA element counts rna_C, rna_N, rna_O, '
     'rna_P, rna_S) derived directly from atomic coordinates.')

# РћђРћђ 3.6 РћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђ
section(doc, 'Machine Learning Modeling', numbered='3.6', level=2)

body(doc,
     'We employ support vector regression [28] with a radial basis function (RBF) kernel',
     first_para=True)
eq(doc,
   'K(x, xРђ▓)  =  exp( Рѕњ╬│ Рђќx Рѕњ xРђ▓Рђќ┬▓ )',
   24)
body(doc,
     'implemented via scikit-learn [29] throughout CAML-RNA. SVR-RBF is well-suited '
     'to datasets of the size considered here (nРђ»=Рђ»143): unlike '
     'gradient-boosted trees or deep neural networks, SVR achieves strong '
     'generalization with a small number of effective parameters and remains '
     'robust to high-dimensional sparse features after standard z-score scaling.')

body(doc,
     'Global Model.  '
     'Hyperparameters CРђ»РѕѕРђ»{0.01, 0.1, 1, 10, 100} and '
     '╬│Рђ»РѕѕРђ»{scale, auto} are selected by 5-fold inner '
     'cross-validation (maximizing R┬▓) inside each outer fold of a '
     '5-fold nested cross-validation scheme. Out-of-fold (OOF) predictions '
     'for all 143 complexes provide an unbiased estimate of generalization '
     'performance.')

body(doc,
     'Subtype-Aware Feature Selection.  '
     'For each structural subtype S with n_S complexes '
     '(4Рђ»РЅцРђ»n_SРђ»РЅцРђ»61), we compare the global OOF '
     'predictions against predictions from a dedicated LOO-CV model trained on '
     'candidate feature set X_S. The override criterion is:')
eq(doc,
   'Apply override if:   R(yРѓЏ, ┼иРѓЏр┤Ир┤║р┤║) > R(yРѓЏ, '
   '┼иРѓЏрХбр┤ИрхЃрхЄрхЃр┤И)',
   25)
body(doc,
     'where R(┬и, ┬и) denotes Pearson correlation. '
     'LOO-CV is used for subtype-specific models because it maximally utilizes '
     'limited training data while providing an unbiased performance estimate. '
     'Feature candidates include: ES+CS Betti curves, CPF, RNA-FM embeddings, '
     'Morgan fingerprints, physicochemical descriptors, and pairwise '
     'concatenations thereof. Subtype-level overrides are applied first; '
     'riboswitch subclass-level refinements are then applied within '
     'the riboswitch subtype.')

body(doc,
     'Sign Correction for Anti-Correlated Subgroups.  '
     'For subgroups where all feature modalities yield RРђ»<Рђ»0 in '
     'dedicated LOO models (SAM/SAH riboswitch), the dominant affinity signal '
     'is anti-correlated with the chosen descriptor polarity. '
     'After sign-correcting the raw predictions by reflection about their '
     'empirical mean:')
eq(doc,
   '╚│рхб  =  2╚│РѓЏ Рѕњ ┼ирхб ,     i Рѕѕ S',
   26)
body(doc,
     'where ╚│РѓЏРђ»=Рђ»NРѓЏРЂ╗┬╣ РѕЉрхбРѕѕРѓЏРђ»┼ирхб, '
     'we recover RРђ»=Рђ»0.248 for SAM/SAH, representing the empirical '
     'ceiling for this subclass.')

# РћђРћђ 3.7 РћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђ
section(doc, 'Evaluation Metrics', numbered='3.7', level=2)

body(doc,
     'Model performance is quantified by the Pearson correlation coefficient:',
     first_para=True)
eq(doc,
   'R  =  РѕЉрхљ(yрхљр┤▒ Рѕњ ╚│р┤▒)(yрхљрхќ Рѕњ ╚│рхќ)'
   '  /  [ РѕџРѕЉрхљ(yрхљр┤▒ Рѕњ ╚│р┤▒)┬▓ '
   '┬и РѕџРѕЉрхљ(yрхљрхќ Рѕњ ╚│рхќ)┬▓ ]',
   27)
body(doc,
     'where yрхљр┤▒ and yрхљрхќ denote experimental and predicted '
     'pKd for complex m, and ╚│р┤▒, ╚│рхќ are their respective '
     'means over NРђ»=Рђ»143 complexes. '
     'The root-mean-square error (RMSE) is')
eq(doc,
   'RMSE  =  Рѕџ( NРЂ╗┬╣ РѕЉрхљРѓїРѓЂрхј (yрхљр┤▒ '
   'Рѕњ yрхљрхќ)┬▓ )',
   28)
body(doc,
     'Spearman rank correlation ¤Ђ measures monotone agreement independently '
     'of the functional form of the affinityРђЊdescriptor relationship. '
     'Uncertainty in R is estimated by bootstrap resampling '
     '(nрхЄрхњрхњрхЌРђ»=Рђ»10,000, seed 42), reporting the '
     '2.5th and 97.5th percentiles as the 95Рђ»% confidence interval.')

# РЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљ
# 4. MODEL INTERPRETABILITY
# РЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљ
section(doc, 'Model Interpretability', numbered='4.', level=1)

body(doc,
     'A key advantage of CAML-RNA over black-box deep learning approaches is '
     'the direct physical interpretability of the persistent Betti curves. '
     'Because each Betti curve is tied to a specific element pair '
     '(╬▒, ╬▓) or category pair (Bрхб, Lр┤╝), the predictive '
     'contribution of each curve can be mapped back to specific types of '
     'RNAРђЊligand contact, enabling chemical attribution of model predictions.',
     first_para=True)

body(doc,
     'Figure 1bРђЊc contrasts the ╬▓0 and ╬▓1 Betti curves for two '
     'representative complexes: a high-affinity SAM-I riboswitch binder '
     '(pKdРђ»=Рђ»10.96, PDB: 3MXH) and a low-affinity purine riboswitch '
     'binder (pKdРђ»=Рђ»2.51, PDB: 4ERJ). '
     'The ╬▓0 curve tracks how isolated RNAРђЊligand atom-type pairs merge '
     'into connected simplicial components as the filtration radius r increases. '
     'For the high-affinity complex, ╬▓0 decays rapidly at short radii '
     '(rРђ»РЅѕРђ»2РђЊ4Рђ»├Ё), indicating that RNA and ligand atoms '
     'of complementary types are in close, dense, and spatially organized contact. '
     'The low-affinity complex shows a slower, more gradual ╬▓0 decay, '
     'reflecting sparser and less geometrically organized contacts. '
     'In StanleyРђЊReisner terms, the ╬▓0 decay rate corresponds to the '
     'rate at which the StanleyРђЊReisner ideal I(╬ћ╩│) acquires '
     'minimal generators as r increases, directly encoding contact density.')

body(doc,
     'The ╬▓1 curve records the birth and death of 1-cycles (loops) in the '
     'bipartite simplicial complex. '
     'High-affinity complexes exhibit a pronounced ╬▓1 peak at intermediate '
     'radii (rРђ»РЅѕРђ»4РђЊ8Рђ»├Ё), arising from ring-like arrangements '
     'of RNAРђЊligand contacts Рђћ an algebraic signature of the enclosing, '
     'complementary cavity formed around the ligand. '
     'Low-affinity complexes produce smaller or later-appearing ╬▓1 peaks, '
     'consistent with incomplete geometric enclosure. '
     'In the StanleyРђЊReisner framework, the ╬▓1 peak corresponds to a '
     '1-dimensional hole in the induced filtration whose algebraic counterpart is '
     'a degree-3 generator of I(╬ћ╩│) that persists over a range '
     '[rрхЄрхб╩│╩░рхЌ, rрхѕрхЃрхЃрхЌ╩░] encoding '
     'the spatial extent of this enclosure.')

body(doc,
     'Element-specific resolution further enables chemical attribution. '
     'For aptamers, the CРђЊN and NРђЊO element pairs dominate '
     'predictive importance (highest Betti curve variance across complexes), '
     'consistent with hydrogen-bonding networks between nucleobase nitrogen '
     'atoms and ligand carbonyl or amino groups as the primary affinity '
     'determinants. For TPP riboswitches, PРђЊO and PРђЊN element pairs '
     'from phosphateРђЊligand contacts contribute most strongly, matching '
     'the known role of magnesium-coordinated phosphate recognition in TPP '
     'binding. This chemical interpretability distinguishes CAML-RNA from '
     'graph neural networks and sequence-based models, which compress '
     'structural information into latent representations that resist '
     'chemical attribution.')

# РЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљ
# 5. CONCLUSIONS
# РЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљ
section(doc, 'Conclusions', numbered='5.', level=1)

body(doc,
     'We have demonstrated that Commutative Algebra Machine Learning Рђћ '
     'specifically persistent StanleyРђЊReisner theory applied to bipartite '
     'RNAРђЊligand simplicial complexes Рђћ provides a mathematically '
     'rigorous and chemically interpretable framework for RNAРђЊligand '
     'binding affinity prediction. '
     'CAML-RNA achieves RРђ»=Рђ»0.7288 on a 143-complex benchmark spanning '
     'seven RNA structural subtypes, outperforming AffiGrapher, RLaffinity, and '
     'RLASIF while providing unprecedented subtype-level resolution: '
     'aptamers (RРђ»=Рђ»0.940), riboswitch family (RРђ»=Рђ»0.771), '
     'and ribosomal A-site (RРђ»=Рђ»0.763).',
     first_para=True)

body(doc,
     'Three lessons emerge for RNA affinity prediction more broadly. '
     'First, the optimal descriptor modality is strongly subtype-dependent: '
     'topological contact pair features excel for aptamers and TPP riboswitches '
     'where geometric complementarity dominates; RNA-FM embeddings prevail '
     'where global fold context is informative (ribosomal A-site, FMN/FAD '
     'riboswitch); and simple physicochemical descriptors outperform complex '
     'fingerprints for ligand-chemotype-constrained subgroups (purine riboswitches). '
     'Second, hard prediction ceilings exist for structurally heterogeneous '
     'subgroups (duplex groove, other/misc, SAM/SAH) where current structural '
     'representations do not encode the dominant affinity determinants. '
     'These ceilings likely reflect contributions from cation coordination, '
     'RNA conformational dynamics, and cryptic binding sites that are not '
     'captured in static three-dimensional coordinates. ')

body(doc,
     'CAML-RNA establishes commutative algebra machine learning as a productive '
     'research direction for nucleic acid drug discovery. '
     'Natural extensions include application to DNAРђЊligand systems, '
     'RNAРђЊprotein interfaces, and integration with structure prediction '
     'pipelines (e.g., RhoFold, AlphaFold-RNA) as high-quality predicted '
     'structures become available for larger training sets. '
     'We make all code, pre-computed features, and out-of-fold prediction arrays '
     'publicly available to facilitate reproducible benchmarking in the '
     'RNA informatics community.')

# РЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљ
# ACKNOWLEDGEMENTS
# РЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљ
section(doc, 'Acknowledgements', level=1)

body(doc,
     'The authors declare no external funding for this work.',
     first_para=True)

# РЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљ
# DATA AVAILABILITY STATEMENT
# РЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљ
section(doc, 'Data Availability Statement', level=1)

body(doc,
     'The 143-complex RNAРђЊligand dataset, all pre-computed feature arrays '
     '(ES+CS Betti curves, CPF, RNA-FM embeddings), out-of-fold prediction '
     'files, trained model pipelines, and all analysis scripts (Steps 1РђЊ33) '
     'are available at https://github.com/[repository]. '
     'The NA-L nucleic acidРђЊligand database from which the dataset derives '
     'is available at http://[na-l-database-url].',
     first_para=True)

# РЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљ
# REFERENCES
# РЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљРЋљ
section(doc, 'References', level=1)

refs = [
    '[1]  Serganov, A.; Nudler, E. A Decade of Riboswitches. '
    'Cell 2013, 152, 17РђЊ24.',

    '[2]  Mortimer, S. A.; Kidwell, M. A.; Doudna, J. A. Insights into RNA '
    'Structure and Function from Genome-Wide Studies. '
    'Nat. Rev. Genet. 2014, 15, 469РђЊ479.',

    '[3]  Connell, S. R. et al. Structural Basis for Interaction of an Antibiotic '
    'Peptide Deformylase Inhibitor with the Ribosomal Peptidyl Transferase Center. '
    'Science 2006, 313, 1020РђЊ1023.',

    '[4]  Thomas, J. R.; Hergenrother, P. J. Targeting RNA with Small Molecules. '
    'Chem. Rev. 2008, 108, 1171РђЊ1224.',

    '[5]  Winkler, W.; Nahvi, A.; Breaker, R. R. Thiamine Derivatives Bind Messenger '
    'RNAs Directly to Regulate Bacterial Gene Expression. '
    'Nature 2002, 419, 952РђЊ956.',

    '[6]  Breaker, R. R. Riboswitches and the RNA World. '
    'Cold Spring Harb. Perspect. Biol. 2012, 4, a003566.',

    '[7]  Donlic, A.; Hargrove, A. E. Targeting RNA in Mammalian Systems with Small '
    'Molecules. Wiley Interdiscip. Rev. RNA 2018, 9, e1477.',

    '[8]  Warner, K. D.; Hajdin, C. E.; Weeks, K. M. Principles for Targeting RNA '
    'with Drug-Like Small Molecules. '
    'Nat. Rev. Drug Discov. 2018, 17, 547РђЊ558.',

    '[9]  Luo, J.; Wei, W.; Waldisph├╝l, J.; Moitessier, N. Challenges and '
    'Current Status of Computational Methods for Docking Small Molecules to Nucleic '
    'Acids. Eur. J. Med. Chem. 2019, 168, 414РђЊ425.',

    '[10]  Chen, L. et al. Computational Prediction of RNAРђЊLigand Interactions: '
    'A Survey and Outlook. '
    'Brief. Bioinform. 2022, 23, bbac085.',

    '[11]  Pfeffer, P.; Gohlke, H. DrugScoreRNA: A Knowledge-Based Scoring Function '
    'to Predict RNAРђЊLigand Interactions. '
    'J. Chem. Inf. Model. 2007, 47, 1868РђЊ1876.',

    '[12]  Bian, Y.; Xia, J. Computational Approaches to RNAРђЊSmall Molecule '
    'Binding: A Review. ACS Omega 2022, 7, 6308РђЊ6314.',

    '[13]  Guillemot, J. C. et al. Scoring Functions for RNAРђЊLigand Binding '
    'Affinity: New Developments and Challenges. '
    'J. Comput.-Aided Mol. Des. 2024, 38, 19.',

    '[14]  [Author(s)]. AffiGrapher: Graph Neural Network-Based RNAРђЊLigand '
    'Binding Affinity Prediction. [Journal] 2023, [vol], [pages].',

    '[15]  [Author(s)]. RLaffinity: Reinforcement Learning-Augmented RNAРђЊLigand '
    'Affinity Scoring. [Journal] 2022, [vol], [pages].',

    '[16]  [Author(s)]. RLASIF: RNAРђЊLigand Binding Affinity Prediction via '
    'Structural Interface Fingerprints. [Journal] 2023, [vol], [pages].',

    '[17]  [Author(s)]. DeepRSMA: Deep Learning for RNAРђЊSmall Molecule Affinity '
    'Prediction. [Journal] 2023, [vol], [pages].',

    '[18]  [Author(s)]. RSAPred: RNA Secondary Structure-Aware Affinity Prediction. '
    '[Journal] 2022, [vol], [pages].',

    '[19]  Edelsbrunner, H.; Letscher, D.; Zomorodian, A. Topological Persistence and Simplification. Discrete Comput. Geom. 2002, 28, 511├б┬ђ┬Њ533.',

    '[20]  Suwayyid, F.; Wei, G.-W. Persistent StanleyРђЊReisner Theory for '
    'Topological Data Analysis. '
    'J. Chem. Theory Comput. 2024, 20, 1248РђЊ1264.',

    '[21]  Miller, E.; Sturmfels, B. Combinatorial Commutative Algebra; '
    'Springer: New York, 2005.',

    '[22]  Bruns, W.; Herzog, J. CohenРђЊMacaulay Rings, revised ed.; '
    'Cambridge University Press: Cambridge, 1998.',

    '[23]  Feng, H.; Suwayyid, F.; Zia, M.; Wee, J.; Hozumi, Y.; Chen, C.-L.; '
    'Wei, G.-W. CAML: Commutative Algebra Machine Learning Рђћ A Case Study on '
    'ProteinРђЊLigand Binding Affinity Prediction. '
    'J. Chem. Inf. Model. 2025, 65, 6732РђЊ6743.',

    '[24]  Edelsbrunner, H.; Letscher, D.; Zomorodian, A. Topological Persistence '
    'and Simplification. Discrete Comput. Geom. 2002, 28, 511РђЊ533.',

    '[25]  Bauer, U. Ripser: Efficient Computation of VietorisРђЊRips Persistence '
    'Barcodes. J. Appl. Comput. Topol. 2021, 5, 391РђЊ423.',

    '[26]  Chen, J. et al. Interpretable RNA Foundation Model from Unannotated Data '
    'for Highly Accurate RNA Structure and Function Predictions. '
    'arXiv 2022, 2204.00300.',

    '[27]  Rogers, D.; Hahn, M. Extended-Connectivity Fingerprints. '
    'J. Chem. Inf. Model. 2010, 50, 742РђЊ754.',

    '[28]  Cortes, C.; Vapnik, V. Support-Vector Networks. '
    'Mach. Learn. 1995, 20, 273РђЊ297.',

    '[29]  Pedregosa, F. et al. Scikit-learn: Machine Learning in Python. '
    'J. Mach. Learn. Res. 2011, 12, 2825РђЊ2830.',

    '[30]  Landrum, G. RDKit: Open-Source Cheminformatics. '
    'https://www.rdkit.org, 2013.',

    '[31]  Rose, P. W. et al. The RCSB Protein Data Bank: Views of Structural '
    'Biology for Basic and Applied Research and Education. '
    'Nucleic Acids Res. 2015, 44, D457РђЊD466.',

    '[32]  Zomorodian, A.; Carlsson, G. Computing Persistent Homology. '
    'Discrete Comput. Geom. 2005, 33, 249РђЊ274.',

    '[33]  Wee, J.; Xia, K. Persistent Spectral Based Machine Learning for '
    'ProteinРђЊLigand Binding Affinity Prediction. '
    'Sci. Adv. 2021, 7, eabf5135.',

    '[34]  Cang, Z.; Wei, G.-W. TopologyNet: Topology Based Deep Convolutional and '
    'Multi-Task Neural Networks for Biomolecular Property Predictions. '
    'PLoS Comput. Biol. 2017, 13, e1005690.',

    '[35]  Kuzmin, M. et al. 3D Graph Neural Networks for RNA Secondary '
    'Structure-Based Ligand Binding Prediction. '
    'J. Chem. Inf. Model. 2021, 61, 1937РђЊ1950.',

    '[36]  Zhou, G. et al. GraphDTA: Predicting DrugРђЊTarget Binding Affinity '
    'with Graph Neural Networks. '
    'Bioinformatics 2020, 37, 1140РђЊ1147.',

    '[37]  [Author(s)]. NA-L: Curated Database of Nucleic AcidРђЊLigand Binding '
    'Affinities. [Journal] 2020, [vol], [pages].',

    '[38]  Rose, P. W. et al. The RCSB Protein Data Bank. '
    'Nucleic Acids Res. 2015, 44, D457РђЊD466.',
]

for ref in refs:
    p = doc.add_paragraph()
    _set_spacing(p, 1, 2, 1.0)
    p.paragraph_format.left_indent        = Inches(0.4)
    p.paragraph_format.first_line_indent  = Inches(-0.4)
    run = p.add_run(ref)
    _set_font(run, size=10)

# РћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђРћђ
doc.save(OUT_PATH)
print(f'Saved: {OUT_PATH}')
